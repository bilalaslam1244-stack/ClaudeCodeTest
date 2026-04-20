import asyncio
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from bot.config import CJK_FONT_PATH, OUTPUT_DIR, CLAUDE_SONNET_MODEL
from bot.services import claude_service

os.makedirs(OUTPUT_DIR, exist_ok=True)


# ── Claude content generation ─────────────────────────────────────────────────

async def generate_content(prompt: str, intent: str = "doc_generate") -> str:
    system = (
        "You are an expert business writer and analyst. "
        "Produce professional, well-structured content based ONLY on data and information the user provides. "
        "Use clear headings and bullet points where appropriate. "
        "CRITICAL: If the user's request contains no actual data, figures, or source material, "
        "respond with exactly: DATA_REQUIRED: <one sentence describing what data is needed>. "
        "Do NOT invent, estimate, or use placeholder numbers or made-up statistics."
    )
    return await asyncio.to_thread(
        claude_service.chat_with_intent, intent, system, prompt, 4096
    )


async def generate_meeting_minutes(transcript: str) -> str:
    system = (
        "You are an expert minute-taker. Given a meeting transcript, produce "
        "structured Minutes of Meeting (MOM) with exactly these sections using '## ' headings:\n\n"
        "## Meeting Details\n"
        "Date, time, location (if mentioned), and participants.\n\n"
        "## Agenda / Topics Discussed\n"
        "Bullet points of main topics covered.\n\n"
        "## Key Decisions\n"
        "Bullet points of decisions reached.\n\n"
        "## Action Items\n"
        "Each item on its own line: - [Person] — [Task] — [Deadline if mentioned]\n\n"
        "## Summary\n"
        "2-3 sentence high-level summary of the meeting.\n\n"
        "## Full Transcript\n"
        "The verbatim transcript.\n\n"
        "IMPORTANT formatting rules:\n"
        "- Use '## ' for section headings only\n"
        "- Use '- ' for bullet points\n"
        "- Do NOT use ** or * or _ for emphasis\n"
        "- Do NOT use markdown bold/italic anywhere\n"
        "- Be concise and professional"
    )
    return await asyncio.to_thread(
        claude_service.chat_with_intent,
        "meeting_minutes",
        system,
        f"TRANSCRIPT:\n{transcript}",
        8192,
    )


# ── Document writers ───────────────────────────────────────────────────────────

def _write_docx(content: str, title: str) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # Body — split on lines, detect headings (##)
    import re as _re
    def _clean(s):
        s = _re.sub(r'\*{1,3}', '', s)
        s = _re.sub(r'_{1,2}', '', s)
        return s.strip()

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(_clean(stripped[3:]), level=2)
        elif stripped.startswith("### "):
            p = doc.add_heading(_clean(stripped[4:]), level=3)
        elif stripped.startswith("• ") or stripped.startswith("- "):
            p = doc.add_paragraph(_clean(stripped[2:]), style="List Bullet")
        else:
            p = doc.add_paragraph(_clean(stripped))

        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(11)
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            except Exception:
                pass

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title)[:40]
    out_path = os.path.join(OUTPUT_DIR, f"{safe_title}_{timestamp}.docx")
    doc.save(out_path)
    return out_path


def _write_pdf(content: str, title: str) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Register CJK font if available
    font_name = "Helvetica"
    if os.path.exists(CJK_FONT_PATH):
        pdfmetrics.registerFont(TTFont("NotoSansCJK", CJK_FONT_PATH))
        font_name = "NotoSansCJK"

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title)[:40]
    out_path = os.path.join(OUTPUT_DIR, f"{safe_title}_{timestamp}.pdf")

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", fontName=font_name, fontSize=16,
                                 spaceAfter=12, leading=20)
    h2_style = ParagraphStyle("H2", fontName=font_name, fontSize=13,
                              spaceAfter=8, spaceBefore=12, leading=16)
    body_style = ParagraphStyle("Body", fontName=font_name, fontSize=11,
                                leading=16, spaceAfter=6)

    story = [Paragraph(title, title_style), Spacer(1, 0.5*cm)]

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.3*cm))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], h2_style))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], h2_style))
        elif stripped.startswith(("• ", "- ")):
            story.append(Paragraph(f"• {stripped[2:]}", body_style))
        else:
            story.append(Paragraph(stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), body_style))

    doc.build(story)
    return out_path


async def create_document(content: str, title: str, fmt: str = "docx") -> str:
    """Generate a .docx or .pdf file. Returns the file path."""
    if fmt == "pdf":
        return await asyncio.to_thread(_write_pdf, content, title)
    return await asyncio.to_thread(_write_docx, content, title)


def cleanup_output_file(path: str) -> None:
    try:
        if os.path.exists(path):
            os.unlink(path)
    except OSError:
        pass
